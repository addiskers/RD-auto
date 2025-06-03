from flask import Flask, render_template, request, send_file, redirect, url_for
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
import tempfile
import os
from docx.shared import RGBColor

def transform_market_data(data, market_name):
    segments = {}
    current_segment = None
    current_level = 0

    for item, level in data:
        if level not in [0, 1, 2]:
            continue
        if level == 0:
            if f"Global {market_name} Market Size by" in item:
                segment_name = item.replace(f"Global {market_name} Market Size by ", "").split(" & ")[0]
                current_segment = segment_name
                segments[current_segment] = []
            continue

        if current_segment and level > 0:
            if item != "Market Overview":
                if level == 1:
                    segments[current_segment].append(item)
                elif level == 2:
                    if segments[current_segment] and isinstance(segments[current_segment][-1], list):
                        segments[current_segment][-1].append(item)
                    else:
                        segments[current_segment].append([item])
    
    formatted_output = []
    for segment, sub_segments in segments.items():
        formatted_subs = []
        for sub in sub_segments:
            if isinstance(sub, list):
                formatted_subs[-1] += f" ({', '.join(sub)})"
            else:
                formatted_subs.append(sub)

        segment_line = f"Segment {segment}: Sub-Segments {', '.join(formatted_subs)}"
        formatted_output.append(segment_line)
    return formatted_output, segments

def generate_segmental_analysis(segments_data, market_name):
    text = f"Global {market_name} Market is segmented by "
    segment_names = []
    segment_details = []

    for segment, sub_segments in segments_data.items():
        segment_names.append(segment)
        sub_details = []
        for sub in sub_segments:
            if isinstance(sub, list):
                continue
            else:
                sub_details.append(sub)

        if not sub_details:
            segment_details.append(
                f"Based on {segment}, no specific sub-segments were identified."
            )
        elif len(sub_details) > 1:
            joined_sub_details = ", ".join(sub_details[:-1]) + " and " + sub_details[-1]
        else:
            joined_sub_details = sub_details[0]

        if sub_details:
            segment_details.append(
                f"Based on {segment}, the market is segmented into {joined_sub_details}."
            )

    if len(segment_names) > 1:
        text += ", ".join(segment_names) + " and region. "
    else:
        text += segment_names[0] + " and region. "

    text += " ".join(segment_details)
    text = (
        text
        + " Based on region, the market is segmented into North America, Europe, Asia Pacific, Latin America and Middle East & Africa. "
    )
    return text

def title_h1(segments_data, market_name):
    segment_all = []
    
    for segment, sub_segments in segments_data.items():
        top_level_subsegments = []  
        
        for sub in sub_segments:
            if isinstance(sub, str):  
                top_level_subsegments.append(sub)
        if top_level_subsegments:
            subsegments_text = f" ({', '.join(top_level_subsegments[:2])})"
        else:
            subsegments_text = ""

        segment_all.append("By " + segment + subsegments_text)

    if len(segment_all) >= 1:
        text_seg = ", ".join(segment_all)
        text = f"{market_name} Market Size, Share, Growth Analysis, {text_seg}, By Region - Industry Forecast 2025-2032"
    else:
        text = f"{market_name} Market Size, Share, Growth Analysis, By Region - Industry Forecast 2025-2032"

    return text


def export_to_word(data, market_name, value_2023, currency, cagr, companies, output_path="Market_Report.docx"):
    value_2024 = value_2023 * (1 + cagr / 100) ** 1
    value_2032 = value_2023 * (1 + cagr / 100) ** 9
    value_2024 = round(value_2024, 2)
    value_2032 = round(value_2032, 2)

    doc = Document()
    formatted_output, segments = transform_market_data(data, market_name)

    def set_poppins_style(paragraph, size=12, bold=False, color=RGBColor(0, 0, 0)):
        run = paragraph.add_run()
        run.font.name = "Poppins"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Poppins")
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color  
        return run

    title = doc.add_heading(level=1)
    title_run = set_poppins_style(title, size=16, bold=True, color=RGBColor(0, 0, 0))
    title_run.text = "Report Name"
    text_paragraph = doc.add_paragraph()
    text_run = set_poppins_style(text_paragraph, size=12, color=RGBColor(0, 0, 0))
    text_run.text = f"{market_name} Market"

    upcoming = doc.add_heading(level=1)
    upcoming_run = set_poppins_style(upcoming, size=16, bold=True, color=RGBColor(0, 0, 0))
    upcoming_run.text = "Upcoming"
    text_paragraph = doc.add_paragraph()
    text_run = set_poppins_style(text_paragraph, size=12, color=RGBColor(0, 0, 0))
    text_run.text = "No"

    segments_heading = doc.add_heading(level=1)
    segments_heading_run = set_poppins_style(segments_heading, size=16, bold=True, color=RGBColor(0, 0, 0))
    segments_heading_run.text = "Segments"

    for line in formatted_output:
        if line.startswith("Segment"):
            segment_heading = doc.add_heading(level=2)
            segment_run = set_poppins_style(segment_heading, size=16, bold=True, color=RGBColor(0, 0, 0))
            segment_run.text = "Segment"

            segment_name_paragraph = doc.add_paragraph()
            segment_name_run = set_poppins_style(segment_name_paragraph, size=12, color=RGBColor(0, 0, 0))
            segment_name_run.text = line.split(":")[0].replace("Segment", "").strip()

            sub_segment_label = doc.add_heading(level=2)
            sub_segment_label_run = set_poppins_style(
                sub_segment_label, size=16, bold=False, color=RGBColor(0, 0, 0)
            )
            sub_segment_label_run.text = "Sub-Segments"

            sub_segment_paragraph = doc.add_paragraph()
            sub_segment_run = set_poppins_style(sub_segment_paragraph, size=12, color=RGBColor(0, 0, 0))
            sub_segment_run.text = line.split("Sub-Segments")[1].strip()

    market_heading = doc.add_heading(level=1)
    market_heading_run = set_poppins_style(market_heading, size=16, bold=True, color=RGBColor(0, 0, 0))
    market_heading_run.text = "Market Insights"

    text_paragraph = doc.add_paragraph()
    text_run = set_poppins_style(text_paragraph, size=12, color=RGBColor(0, 0, 0))
    text_run.text = (
        f"Global {market_name} Market size was valued at USD {value_2023} {currency} in 2023 "
        f"and is poised to grow from USD {value_2024} {currency} in 2024 to USD {value_2032} {currency} by 2032, "
        f"growing at a CAGR of {cagr}% during the forecast period (2025-2032)."
    )

    market_heading_1 = doc.add_heading(level=1)
    market_heading_run_1 = set_poppins_style(market_heading_1, size=16, bold=True, color=RGBColor(0, 0, 0))
    market_heading_run_1.text = "Segmental Analysis"

    text_paragraph = doc.add_paragraph()
    text_run = set_poppins_style(text_paragraph, size=12, color=RGBColor(0, 0, 0))
    text_run.text = generate_segmental_analysis(segments, market_name)

    top_players_heading = doc.add_paragraph()
    top_players_run = set_poppins_style(top_players_heading, size=12, bold=True, color=RGBColor(0, 0, 0))
    top_players_run.text = "Top Player's Company Profiles"
    
    for c1 in companies.splitlines():
        company_paragraph = doc.add_paragraph(style="List Bullet")
        company_run = set_poppins_style(company_paragraph, size=12, color=RGBColor(0, 0, 0))
        company_run.text = c1.strip()

    H1_text = doc.add_heading(level=1)
    H1_run = set_poppins_style(H1_text, size=16, bold=True, color=RGBColor(0, 0, 0))
    H1_run.text = "H1 Title"

    text_paragraph = doc.add_paragraph()
    title_text = title_h1(segments, market_name)
    text_run = set_poppins_style(
        text_paragraph,
        size=12,
        color=RGBColor(255, 0, 0) if len(title_text.split()) > 35 else RGBColor(0, 0, 0)
    )
    text_run.text = title_text
    doc.save(output_path)
    return output_path

app = Flask(__name__)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
doc_path = os.path.join(BASE_DIR, 'toc.docx')
rd_path = os.path.join(BASE_DIR, 'rd.docx')

if not os.path.exists(doc_path):
    doc = Document()
    doc.save(doc_path)

def get_level(i):
    """Determine level based on indentation style."""
    return i.split(" ", 1)[0].count(".")

def clean(name):
    """Clean and title case the name."""
    a = name.split(" ", 1)[1]
    if "(Page No." in a:
        a = a.split(" (Page No.", 1)[0].strip()
    return a.strip()

def get_level1(i):
    try:
        i = i.replace("\t", " ")
        return i.split(" ", 1)[0].count(".")
    except:
        return i.split(" ", 1)[0].count(".")

def clean1(name):
    try:
        name = name.replace("\t", " ")
        a = name.split(" ", 1)[1]
    except:
        a = name.split(" ", 1)[1]
    if "(Page No." in a:
        a = a.split(" (Page No.", 1)[0].strip()
    return a.strip()

def add_bullet_point_text(doc, text, level):
    paragraph = doc.add_paragraph(text)
    paragraph.style = 'List Paragraph'
    numbering = paragraph._element.get_or_add_pPr().get_or_add_numPr()
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numbering.append(numId)
    numbering.append(ilvl)
    run = paragraph.runs[0]
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    if level == 0:
        run.bold = True 
    paragraph.paragraph_format.line_spacing = 1.5

@app.route("/", methods=["GET", "POST"])
def index():
    data = {
        "kmi": [],
        "toc_entries": [],
        "segments": [],
        "regions": [],
        "companies": [],
        "conclusion":[("Conclusion & Analyst Recommendations",0)]
    }
    raw_segments = []
    market_details = {
        "market_name": request.form.get("market_name", "").strip(),
        "value_2023": float(request.form.get("value_2023") or  0),
        "currency": request.form.get("currency", "million").strip(),
        "cagr": float(request.form.get("cagr") or 0),
    }

    if request.method == "POST":
        market_name = request.form.get("market_name", "").strip()
        if not market_name:
            return render_template("index.html", error="Market name is required!")

        kmi_data = request.form.get("kmi_data", "").strip()
        if kmi_data:
            data["kmi"].extend([(km.title().strip(), 1) for km in kmi_data.splitlines() if km.strip()])

        headings = request.form.getlist("headings[]")
        levels = request.form.getlist("levels[]")
        for heading, level in zip(headings, levels):
            level = int(level)
            if level == 0:
                raw_segments.append(heading.title())
                toc_heading = f"Global {market_name} Market Size by {heading.title()} (2025-2032)"
                data["toc_entries"].append((toc_heading, level))
                data["toc_entries"].append(("Market Overview", 1))
            else:
                data["toc_entries"].append((heading.title(), level))

        segment_data = request.form.get("segment_data", "").strip()
        if segment_data:
            for seg in segment_data.splitlines():
                seg_level = get_level1(seg) - 1
                cleaned = clean1(seg)
                if seg_level == 0:
                    raw_segments.append(cleaned)
                    cleaned_1 = f"Global {market_name} Market Size by {cleaned} (2025-2032)"
                    data["segments"].append((cleaned_1, seg_level))
                    data["segments"].append(("Market Overview", 1))
                else:
                    data["segments"].append((cleaned, seg_level))

        company_data = request.form.get("company_data", "").strip()
        if company_data:
            data["companies"].extend([
                ("Competitive Dashboard", 0),
                ("Top 5 Player Comparison", 1),
                ("Market Positioning of Key Players, 2024", 1),
                ("Strategies Adopted by Key Market Players", 1),
                ("Recent Developments in the Market", 1),
                ("Company Market Share Analysis, 2024", 1),
                ("Key Company Profiles", 0),
            ])

            company_lines = [line.strip() for line in company_data.splitlines() if line.strip()]
            for idx, company_name in enumerate(company_lines):
                data["companies"].append((company_name, 1))
                if idx == 0:
                    data["companies"].extend([
                        ("Company Overview", 2),
                        ("Product Portfolio Overview", 2),
                        ("Financial Overview", 2),
                        ("Key Developments", 2),
                    ])
                    data["companies"].append(("__INSERT_ITALIC_AFTER_KEY_DEVELOPMENTS__", 1))

        regions = [
            ("North America", ["US", "Canada"]),
            ("Europe", ["Germany", "Spain", "France", "UK", "Italy", "Rest of Europe"]),
            ("Asia Pacific", ["China", "India", "Japan", "South Korea", "Rest of Asia-Pacific"]),
            ("Latin America", ["Brazil", "Rest of Latin America"]),
            ("Middle East & Africa", ["GCC Countries", "South Africa", "Rest of Middle East & Africa"]),
        ]

        data["regions"].append((f"Global {market_name} Market Size (2025-2032)", 0))
        segment_text_list = [f"By {segment}" for segment in raw_segments] if raw_segments else ["No segments available"]
        segment_text = ", ".join(segment_text_list)
        for region, subregions in regions:
            data["regions"].append((f"{region} ({segment_text})", 1))
            data["regions"].extend([(subregion, 2) for subregion in subregions])

        toc_content = data["kmi"] + data["toc_entries"] + data["segments"] + data["regions"] + data["companies"] + data["conclusion"]

        toc_temp_file_name = f"TOC_{market_name}_Market_SkyQuest.docx"
        rd_temp_file_name = f"RD_{market_name}_SkyQuest.docx"

        toc_temp_file_path = os.path.join(tempfile.gettempdir(), toc_temp_file_name)
        rd_temp_file_path = os.path.join(tempfile.gettempdir(), rd_temp_file_name)

        toc_doc = Document(doc_path)
        for heading, level in toc_content:
            if heading == "__INSERT_ITALIC_AFTER_KEY_DEVELOPMENTS__":
                p = toc_doc.add_paragraph()
                run = p.add_run("Similar information will be covered for below listed companies. Detailed financial will be provided for public listed companies only. List of companies mentioned below are for indicative purpose.")
                run.italic = True
            else:
                add_bullet_point_text(toc_doc, heading, level)

        toc_doc.save(toc_temp_file_path)

        export_to_word(
            data=data["toc_entries"] + data["segments"],
            market_name=market_name,
            value_2023=market_details["value_2023"],
            currency=market_details["currency"],
            cagr=market_details["cagr"],
            companies=company_data,
            output_path=rd_temp_file_path
        )

        return render_template(
            "index.html",
            file_ready=True,
            toc_file_path=toc_temp_file_path,
            rd_file_path=rd_temp_file_path
        )

    return render_template("index.html", file_ready=False)


@app.route("/download")
def download():
    file_path = request.args.get("file_path")
    
    if not file_path or not os.path.exists(file_path):
        return "Error: The file does not exist. Please generate the document first.", 404
    file_name = os.path.basename(file_path)
    
    return send_file(
        file_path,
        as_attachment=True,
        download_name=file_name, 
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5001)
